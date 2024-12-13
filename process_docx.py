import os
import time
from docx import Document


def load_keywords(file_path):
    """Carregar palavras-chave de um arquivo de texto."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return [line.strip() for line in file.readlines()]


def find_and_highlight_keywords(doc_path, keywords, output_folder, reports_folder):
    """Procura palavras-chave em um documento .docx, destaca e salva se encontrar."""
    document = Document(doc_path)
    found = False
    doc_name = os.path.splitext(os.path.basename(doc_path))[0]
    found_keywords = []
    

    # Percorre cada parágrafo procurando palavras-chave
    for paragraph in document.paragraphs:
        for keyword in keywords:
            if keyword.lower() in paragraph.text.lower():
                print(f'Palavra-chave "{keyword}" encontrada em {doc_path}')
                found = True
                if keyword not in found_keywords:
                    found_keywords.append(keyword)
                highlight_keyword(paragraph, keyword)
        
    # Salvar o arquivo modificado na pasta de saída
    if found:
        os.makedirs(output_folder, exist_ok=True)
        output_path = os.path.join(output_folder, os.path.basename(doc_path))
        document.save(output_path)
    
    return doc_name, found_keywords
        

def highlight_keyword(paragraph, keyword):
    """Destaca a palavra-chave encontrada no parágrafo, mantendo a formatação original."""
    new_runs = []

    # Acessa cada run (segmento de texto com a mesma formatação)
    for run in paragraph.runs:
        start = 0
        text = run.text
        while start < len(text):
            idx = text.lower().find(keyword.lower(), start)
            if idx == -1:
                # Se não encontrar a palavra-chave, adiciona o texto restante
                new_runs.append((text[start:], run.font, False))
                break

            if idx > 0:
                # Adiciona o texto anterior à palavra-chave
                new_runs.append((text[start:idx], run.font, False))

            keyword_run = text[idx:idx + len(keyword)]
            # 'True' indica que é uma palavra-chave
            new_runs.append((keyword_run, run.font, True))

            start = idx + len(keyword)

    paragraph.clear()
    for text_part, font, is_keyword in new_runs:
        run = paragraph.add_run(text_part)
        run.font.bold = font.bold
        run.font.italic = font.italic
        run.font.size = font.size
        run.font.color.rgb = font.color.rgb
        run.font.highlight_color = font.highlight_color
        if is_keyword:
            run.font.highlight_color = 6  # Realce em vermelho


def process_folder(keywords_file, input_folder, output_folder, reports_folder):
    """Processa todos os arquivos .docx na pasta de entrada e gera um relatório."""
    try:
        keywords = load_keywords(keywords_file)
    except FileNotFoundError:
        print("Por favor, crie um arquivo keywords.txt contendo cada keyword separada por quebra de linha")
        time.sleep(3)
        return
    
    report_data = []

    for root, _, files in os.walk(input_folder):
        for file in files:
            if file.endswith('.docx'):
                doc_path = os.path.join(root, file)
                doc_name, found_keywords = find_and_highlight_keywords(doc_path, keywords, output_folder, reports_folder)
                if found_keywords:
                    report_data.append((doc_name, found_keywords))

    generate_report(report_data, reports_folder)


def generate_report(report_data, reports_folder):
    os.makedirs(reports_folder, exist_ok=True)
    
    report_number = 1
    while os.path.exists(os.path.join(reports_folder, f'report_{report_number}.txt')):
        report_number += 1
    
    report_path = os.path.join(reports_folder, f'report_{report_number}.txt')

    with open(report_path, 'a', encoding='utf-8') as report:
        for doc_name, keywords in report_data:
            report.write(f'{doc_name}: {", ".join(keywords)}\n')


keywords_file = 'keywords.txt'
input_folder = 'input_docs'
output_folder = 'output_docs'
reports_folder = 'reports'

process_folder(keywords_file, input_folder, output_folder, reports_folder)
